import re
import html2dom
from EasyTable import EasyTable
import docx
from docx.enum.table import WD_ROW_HEIGHT
from docx.shared import Inches
from docx.shared import Cm
from docx.shared import Pt
#from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

pattern = '###.*?###'
r = re.compile(pattern)
DOCKIND={'1':"소프트웨어요구사항명세서",'2':"소프트웨어설계기술서"}
class DocxComposer:
	def __init__(self, cursor, document,documentid, componentid):
		self.cursor=cursor
		self.document=document
		self.documentid=documentid
		self.componentid=componentid

	def AddHeader(self,text1,text2):
		self.document.add_section()
		newSection=self.document.sections[-1]
		newSection.header.is_linked_to_previous=False
		p=newSection.header.paragraphs[0]._element
		p.getparent().remove(p)
		table=newSection.header.add_table(1,2,Inches(7.5))
		table.rows[0].cells[0].width=Inches(4.5)
		table.rows[0].cells[1].width=Inches(3)
		
		p=table.rows[0].cells[0].paragraphs[0]
		self.insertField(p, "CSCI")
		p.add_run().text=DOCKIND[self.documentid]+"("
		self.insertField(p, "버전")
		p.add_run().text=")"

		for r in table.rows[0].cells[0].paragraphs[0].runs:
			r.font.size=Pt(12)
		

		table.rows[0].cells[1].text=text2
		table.rows[0].cells[1].paragraphs[0].runs[0].font.size=Pt(12)
		table.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

		table.rows[0].height_rule = WD_ROW_HEIGHT.EXACTLY
		table.rows[0].height=Cm(0.8)

		for cell in table.rows[0].cells:
			tcPr = cell._tc.tcPr
			tcBorders = OxmlElement('w:tcBorders')
			top = OxmlElement('w:bottom')
			top.set(qn('w:val'), 'single')
			top.set(qn('w:sz'), str(25))
			top.set(qn('w:space'), '0')
			top.set(qn('w:color'), 'auto')
			tcBorders.append(top)
			tcPr.append(tcBorders)

	def insertDOC(self, item ):
		style = str(int(len(item[5])/2))+"단계"
		if style == "0단계":
			style = None
		if(style == "1단계"): #페이지 분리 & 머리글 달기
			self.AddHeader("#CSCI# 소프트웨어요구사항명세서(V#버전#)",item[6])
			print(style)
		self.document.add_paragraph(item[6], style)
		self.insertDescription(item[9])

	def insertSRSReq(self, item):
		Etable = EasyTable(self.document, rows=4, cols=2, style="a6")
		#tbl = table._tbl  # get xml element in table
		Etable.TableAlign("center")
		Etable.TableColumnsWidth([1.2, 4])
		Etable.TableCellText(["식별자", item[5], "요구사항", item[6], "출처", item[7], "관련 USE Case", item[8]])
		Etable.TableCellBGColor([(0, 0), (1, 0), (2, 0), (3, 0)])
		Etable.TableBordersColumns(1, 10, "left")
		Etable.TableBorders(10)
		self.document.add_paragraph()
		
		self.insertDescription( item[9])
		
	def insertSDDReq(self,  item):
		Etable = EasyTable(self.document, rows=4, cols=2, style="a6")
		#tbl = table._tbl  # get xml element in table
		Etable.TableAlign("center")
		Etable.TableColumnsWidth([1.2, 4])
		Etable.TableCellText(["식별자", item[5], "요구사항", item[6], "출처", item[7], "관련 USE Case", item[8]])
		Etable.TableCellBGColor([(0, 0), (1, 0), (2, 0), (3, 0)])
		Etable.TableBordersColumns(1, 10, "left")
		Etable.TableBorders(10)
		self.document.add_paragraph()
		self.insertDescription( item[9])
		
	def insertSTDReq(self, item):
		self.document.add_paragraph(item[5]+"("+item[6]+")", "2단계")
		self.document.add_paragraph("요구사항", "3단계")
		self.macro6(item)
		self.insertDescription(item[9])

	def insertDescription(self, html):
		count = 0
		if(html==None):
			html="<p></p>"
		for item in html2dom.html2dom(html):
			if(item[0] == 'P'):
				if("#2#") in item[2]:
					self.document.add_paragraph(item[2][3:], "2단계")
					count=0
					continue
				if("#3#") in item[2]:
					self.document.add_paragraph(item[2][3:], "3단계")
					count=0
					continue
				if("#4#") in item[2]:
					self.document.add_paragraph(item[2][3:], "4단계")
					count=0
					continue
				if("###") in item[2]:
					self.insertFieldCodeP( item[2])
					count=0
				elif("#MACRO#") in item[2]:
					self.macro( item[2])
					count=0
				elif("그림#" in item[2]):
					self.insertPictureFiled( item[2].split("그림#")[1])
					count=0
				elif("표#" in item[2]):
					self.insertTableFiled( item[2].split("표#")[1])
					count=0
				else:
					p = self.document.add_paragraph(item[2])
					if(self.document.paragraphs[-1].text == ""):
						count+=1
					try:
						if "center" in item[1]["style"]:
							p.alignment = WD_ALIGN_PARAGRAPH.CENTER
						if "right" in item[1]["style"]:
							p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
					except:
						None

			if(item[0] == 'img'):
				count=0
				width = "640"
				try:
					width = int(item[1]['style'].split(
						"width:")[1].split("px;")[0])
				except:
					None
				try:
					self.document.add_picture("./uploads/"+item[1]['src'].split("uploads/")[1], width=Inches(width/100))
					if("fr-fic" in item[1]["class"]):
						last_paragraph = self.document.paragraphs[-1]
						last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
				except:
					print("img file not found-./uploads/"+item[1]['src'].split("uploads/")[1])
					self.document.add_picture("./plugins/template/lostimage.png", width=Inches(width/100))
					last_paragraph = self.document.paragraphs[-1]
					last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

			if(item[0] == 'table'):
				try:
					count=0
					row = len(item[1])
					consplancount = 0
					for cell in item[1][0]:
						if "colspan" in cell[1]:
							consplancount += int(cell[1]["colspan"])-1
					col = len(item[1][0])+consplancount
					width = 100
					if("style" in item[2] and "width" in item[2]["style"]):
						width = int(item[2]["style"].split("width: ")[1].split("%")[0])
					etable = EasyTable(self.document, row, col, style="a7")
					etable.TableAlign("center")
					etable.TableBorders(10)
					etable.TableSpan(item[1])
					etable.TableProperty(item[1], 6*width/100)
					self.document.add_paragraph()
				except:
					None

		
			if(count>3):
				self.document.add_page_break()


	def insertUC(self, item):
		p = self.document.add_paragraph('A plain paragraph having some ')
		p.add_run('bold').bold = True
		p.add_run(' and some ')
		p.add_run('italic.').italic = True
		self.document.add_page_break()



	def insertFieldCodeP(self, text):
		p = self.document.add_paragraph()
		result = r.findall(text)
		pos = 0
		for t in result:
			new_pos = text.find(t, pos)
			if(new_pos-pos > 0):
				p.add_run(text[pos:new_pos])
			self.insertField(p, t.replace("#", ""))
			pos = new_pos+len(t)
		p.add_run(text[pos:])


	def insertField(self, p, FieldText):
		r2 = p.add_run()
		r3 = p.add_run()
		r4 = p.add_run()
		#<w:fldChar w:fldCharType="begin"/>
		fldChar1 = docx.oxml.shared.OxmlElement('w:fldChar')
		fldChar1.set(docx.oxml.ns.qn('w:fldCharType'), "begin")
		fldChar2 = docx.oxml.shared.OxmlElement('w:instrText')
		fldChar2.text = "DOCPROPERTY  "+FieldText+"  \* MERGEFORMAT"
		fldChar2.set(docx.oxml.ns.qn('xml:space'), 'preserve')
		fldChar3 = docx.oxml.shared.OxmlElement('w:fldChar')
		fldChar3.set(docx.oxml.ns.qn('w:fldCharType'), "end")
		r2._r.append(fldChar1)
		r3._r.append(fldChar2)
		r4._r.append(fldChar3)


	def insertPictureFiled(self, text):
		p = self.document.add_paragraph("그림 ","그림제목")
		p.alignment = WD_ALIGN_PARAGRAPH.CENTER
		#p._p.addnext(Seq_Picture)
		#r1 = p.add_run()
		r2 = p.add_run()
		r3 = p.add_run()
		r4 = p.add_run()
		r5 = p.add_run()
		#r1.text = "그림 "
		r5.text = " "+text
		#<w:fldChar w:fldCharType="begin"/>
		fldChar1 = docx.oxml.shared.OxmlElement('w:fldChar')
		fldChar1.set(docx.oxml.ns.qn('w:fldCharType'), "begin")
		fldChar2 = docx.oxml.shared.OxmlElement('w:instrText')
		fldChar2.text = "SEQ Table \* MERGEFORMAT"
		fldChar2.set(docx.oxml.ns.qn('xml:space'), 'preserve')
		fldChar3 = docx.oxml.shared.OxmlElement('w:fldChar')
		fldChar3.set(docx.oxml.ns.qn('w:fldCharType'), "end")
		r2._r.append(fldChar1)
		r3._r.append(fldChar2)
		r4._r.append(fldChar3)


	def insertTableFiled(self, text):
		p = self.document.add_paragraph("표 ","표제목")
		p.alignment = WD_ALIGN_PARAGRAPH.CENTER
		#p._p.addnext(Seq_Picture)
		#r1 = p.add_run()
		r2 = p.add_run()
		r3 = p.add_run()
		r4 = p.add_run()
		r5 = p.add_run()
		#r1.text = "표 "
		r5.text = " "+text
		#<w:fldChar w:fldCharType="begin"/>
		fldChar1 = docx.oxml.shared.OxmlElement('w:fldChar')
		fldChar1.set(docx.oxml.ns.qn('w:fldCharType'), "begin")
		fldChar2 = docx.oxml.shared.OxmlElement('w:instrText')
		fldChar2.text = "SEQ Picture \* MERGEFORMAT"
		fldChar2.set(docx.oxml.ns.qn('xml:space'), 'preserve')
		fldChar3 = docx.oxml.shared.OxmlElement('w:fldChar')
		fldChar3.set(docx.oxml.ns.qn('w:fldCharType'), "end")
		r2._r.append(fldChar1)
		r3._r.append(fldChar2)
		r4._r.append(fldChar3)

	def macro(self, macroID):
		print(macroID)
		print(macroID.split("#MACRO#")[1])
		fp=eval("self.macro"+macroID.split("#MACRO#")[1])
		
		fp()
		
	def macro1(self):
		query = "Select * from items where document="+self.documentid+" and component="+self.componentid+" and type=11 Order by ordering asc;"
		self.cursor.execute(query)
		result=self.cursor.fetchall()
		text = ["요구사항 식별자", "요구사항 식별자명", "관련항목"]

		for row in result:
			text.append(row[5])
			try:
				text.append(row[6].split("[")[1].split("]")[0])
			except:
				text.append(row[6])
			text.append(row[7])

		etable = EasyTable(self.document, len(result)+1, 3, style="a7")
		etable.TableAlign("center")
		etable.TableBorders(10)
		etable.TableColumnsWidth([1.2, 2.5, 1.9])
		etable.TableCellText(text)
		etable.TableCellBGColor([(0, 0), (0, 1), (0, 2)])
		self.document.add_paragraph()
		
	def macro2(self):
		query = "Select * from items where document="+self.documentid+" and component="+self.componentid+" and type=11 Order by ordering asc;"
		self.cursor.execute(query)
		result=self.cursor.fetchall()
		text = ["요구사항 식별자", "데모","시험","분석","검사","특수","내용"]

		for row in result:
			text.append(row[5])
			for testtype in range(5):
				if(str(testtype)==row[10]):
					text.append("O")
				else:
					text.append("")
			text.append(row[6].split("]")[0]+"]")
		self.insertTableFiled("자격부여방법")
		etable = EasyTable(self.document, len(result)+1, 7, style="a7")
		etable.TableAlign("center")
		etable.TableBorders(10)
		etable.TableColumnsWidth([2.0, 0.3, 0.3, 0.3, 0.3, 0.3, 1.9])
		etable.TableCellText(text)
		etable.TableCellBGColor([(0, 0), (0, 1), (0, 2),(0, 3),(0, 4),(0, 5),(0, 6)])
		self.document.add_paragraph()

		
	def macro3(self):
		query = "Select * from traceability;"
		self.cursor.execute(query)
		links=self.cursor.fetchall()
		
		query = "Select * from items where document="+self.documentid+" and component="+self.componentid+" and type=11 Order by ordering asc;"
		self.cursor.execute(query)
		items=self.cursor.fetchall()

		query = "Select * from items where document=0 and type=1 Order by ordering asc;"
		self.cursor.execute(query)
		SSS=dict()
		for sss in self.cursor.fetchall():
			SSS[sss[0]]=sss[1:]

		query = "Select * from items where document=2 and component="+self.componentid+" and type=21 Order by ordering asc;"
		self.cursor.execute(query)
		SDD=dict()
		for sdd in self.cursor.fetchall():
			SDD[sdd[0]]=sdd[1:]
		
		query = "Select * from items where document=3 and component="+self.componentid+" and type=31 Order by ordering asc;"
		self.cursor.execute(query)
		STD=dict()
		for std in self.cursor.fetchall():
			STD[std[0]]=std[1:]

		

		
		text = ["SW 요구사항 식별자", "체계 요구사항 명세서 식별자", "SW 설계 기술서 식별자","SW 시험 절차서 식별자"]
		
		for row in items:
			#SRS ID
			text.append(row[5])

			#SSS ID
			tempstr=""
			for link in links:
				if(link[0]==1 and link[1]==row[0]):
					sss=SSS.get(link[2],None)
					if(sss == None):
						None
					else:
						tempstr+=sss[4]+"		"
			if(tempstr==""):
				tempstr="추후결정"
			text.append(tempstr)

			#SDD ID
			tempstr=""
			for link in links:
				if(link[0]==2 and link[2]==row[0]):
					sdd=SDD.get(link[1],None)
					if(sdd == None):
						None
					else:
						tempstr+=sdd[4]+"		"
			if(tempstr==""):
				tempstr="추후결정"
			text.append(tempstr)

			#STD ID
			tempstr=""
			for link in links:
				if(link[0]==3 and link[2]==row[0]):
					std=STD.get(link[1],None)
					if(std == None):
						None
					else:
						tempstr+=std[4]+"		"
			if(tempstr==""):
				tempstr="추후결정"
			text.append(tempstr)

		self.insertTableFiled("요구사항추적성")
		etable = EasyTable(self.document, len(items)+1, 4, style="a7")
		etable.TableAlign("center")
		etable.TableBorders(10)
		etable.TableColumnsWidth([1.4, 1.4, 1.4, 1.4])
		etable.TableCellText(text)
		etable.TableCellBGColor([(0, 0), (0, 1), (0, 2), (0, 3)])
		self.document.add_paragraph()

	
	def macro4(self):
		query = "Select * from traceability;"
		self.cursor.execute(query)
		links=self.cursor.fetchall()
		
		query = "Select * from items where document="+self.documentid+" and component="+self.componentid+" and type=21 Order by ordering asc;"
		self.cursor.execute(query)
		items=self.cursor.fetchall()

		
		query = "Select * from items where document=1 and component="+self.componentid+" and type=11 Order by ordering asc;"
		self.cursor.execute(query)
		SRS=dict()
		for srs in self.cursor.fetchall():
			SRS[srs[0]]=srs[1:]
		
		query = "Select * from items where document=3 and component="+self.componentid+" and type=31 Order by ordering asc;"
		self.cursor.execute(query)
		STD=dict()
		for std in self.cursor.fetchall():
			STD[std[0]]=std[1:]
		
		text = ["SW 설계 기술서 식별자", "SW 요구사항 식별자", "SW 시험절차서 식별자"]
		
		for row in items:
			#SDD ID
			text.append(row[5])

			#SRS ID
			SRSs=[]
			tempstr=""
			for link in links:
				if(link[0]==2 and link[1]==row[0]):
					srs=SRS.get(link[2],None)
					if(srs == None):
						None
					else:
						SRSs.append(link[2])
						tempstr+=srs[4]+"		"
			if(tempstr==""):
				tempstr="추후결정"
			text.append(tempstr)
			#STD ID
			tempstr=""
			for link in links:
				if(link[0]==3 and link[2] in SRSs):
					std=STD.get(link[1],None)
					if(std == None):
						None
					else:
						tempstr+=std[4]+"		"
			if(tempstr==""):
				tempstr="추후결정"
			text.append(tempstr)

		self.insertTableFiled("요구사항추적성")
		etable = EasyTable(self.document, len(items)+1, 3, style="a7")
		etable.TableAlign("center")
		etable.TableBorders(10)
		etable.TableColumnsWidth([1.8, 1.8, 1.8 ])
		etable.TableCellText(text)
		etable.TableCellBGColor([(0, 0), (0, 1), (0, 2)])
		self.document.add_paragraph()

	def macro5(self):
		query = "Select * from traceability;"
		self.cursor.execute(query)
		links=self.cursor.fetchall()
		
		query = "Select * from items where document="+self.documentid+" and component="+self.componentid+" and type=31 Order by ordering asc;"
		self.cursor.execute(query)
		items=self.cursor.fetchall()

		
		query = "Select * from items where document=1 and component="+self.componentid+" and type=11 Order by ordering asc;"
		self.cursor.execute(query)
		SRS=dict()
		for srs in self.cursor:
			SRS[srs[0]]=srs[1:]
		
		query = "Select * from items where document=3 and component="+self.componentid+" and type=21 Order by ordering asc;"
		self.cursor.execute(query)
		SDD=dict()
		for sdd in self.cursor:
			SDD[sdd[0]]=sdd[1:]
		
		text = ["시험 식별자", "시험항목", "시험 명세서 항목"]
		count=1
		for row in items:
			#STD ID
			count+=1
			text.append(row[5])
			text.append(row[6])
			text.append("5."+str(count))

		self.insertTableFiled("시험절차서 식별자 목록")
		etable = EasyTable(self.document, len(items)+1, 3, style="a7")
		etable.TableAlign("center")
		etable.TableBorders(10)
		etable.TableColumnsWidth([1.6, 2.3, 1.6 ])
		etable.TableCellText(text)
		etable.TableCellBGColor([(0, 0), (0, 1), (0, 2)])
		self.document.add_paragraph()

	def macro6(self, item):
		query = "Select link2 from traceability where type=3 and link1="+str(item[0])+";"
		self.cursor.execute(query)
		links=[]
		for link in self.cursor.fetchall():
			links.append(str(link[0]))
		print(",".join(links))
		
		query = "Select * from items where document=1 and type=11 and uid in ("+",".join(links)+") Order by ordering asc;"
		self.cursor.execute(query)
		items=self.cursor.fetchall()
		print(items)
		text = []
		count=1
		text.append("시험식별자")
		text.append("\n".join([a[5] for a in items ]))
		text.append("요구사항")
		text.append("\n".join([a[5] for a in items ]))
		print()
		etable = EasyTable(self.document, 2, 2, style="a7")
		etable.TableAlign("center")
		etable.TableBorders(10)
		etable.TableColumnsWidth([1.3, 4.6])
		etable.TableCellText(text)
		etable.TableCellBGColor([(0, 0), (1, 0)])
		self.document.add_paragraph()

	def macro7(self):
		query = "Select * from traceability;"
		self.cursor.execute(query)
		links=self.cursor.fetchall()

		query = "Select * from items where document=3 and component="+self.componentid+" and type=31 Order by ordering asc;"
		self.cursor.execute(query)
		items=self.cursor.fetchall()

		query = "Select * from items where document=1 and component="+self.componentid+" and type=11 Order by ordering asc;"
		self.cursor.execute(query)
		SRS=dict()
		for srs in self.cursor.fetchall():
			SRS[srs[0]]=srs[1:]

		query = "Select * from items where document=2 and component="+self.componentid+" and type=21 Order by ordering asc;"
		self.cursor.execute(query)
		SDD=dict()
		for sdd in self.cursor.fetchall():
			SDD[sdd[0]]=sdd[1:]

		text = ["SW 시험절차서 식별자", "SW 요구사항 명세서 식별자", "SW 설계기술서 식별자"]
		for row in items:
			#SDD ID
			text.append(row[5])

			#SRS ID
			SRSs=[]
			tempstr=""
			for link in links:
				if(link[0]==2 and link[1]==row[0]):
					srs=SRS.get(link[2],None)
					if(srs == None):
						None
					else:
						SRSs.append(link[2])
						tempstr+=srs[4]+"		"
			if(tempstr==""):
				tempstr="추후결정"
			text.append(tempstr)

			#STD ID
			tempstr=""
			for link in links:
				if(link[0]==3 and link[2] in SRSs):
					sdd=SDD.get(link[1],None)
					if(sdd == None):
						None
					else:
						tempstr+=sdd[4]+"		"
			if(tempstr==""):
				tempstr="추후결정"
			text.append(tempstr)

		self.insertTableFiled("요구사항추적성")
		etable = EasyTable(self.document, len(items)+1, 3, style="a7")
		etable.TableAlign("center")
		etable.TableBorders(10)
		etable.TableColumnsWidth([1.8, 2.2, 1.8 ])
		etable.TableCellText(text)
		etable.TableCellBGColor([(0, 0), (0, 1), (0, 2)])
		self.document.add_paragraph()

	#Macro Table
	#1: SRS 표 요구사항 목록
	#2: SRS 표 개발시험평가방법
	#3: SRS 표 요구사항추적성
	#4: SDD 표 요구사항추적성
	#5: STD 표 시험항목 목록
	#6: STD 표 시험 요구사항
	#7: STD 표 요구사항 추적성
	#8: SRS 4차 인터페이스
