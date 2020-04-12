import docx
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx import Document
#from docx.oxml.table import CT_Row, CT_Tc
from docx.shared import Inches
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH


class EasyTable:
    def __init__(self, Document, rows, cols, style):
        self.table = Document.add_table(rows, cols)

    def TableCellText(self, textList):
        for rows in self.table.rows:
            for cell in rows.cells:
                text= (textList.pop(0) if len(textList) > 0 else "")
                cell.text = "" if text ==None else text
                #cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def TableSpan(self, item):
        spanlist = []
        for rc, row in enumerate(item):
            for cc, cell in enumerate(row):
                if("rowspan" in cell[1] and "colspan" in cell[1]):
                    spanlist.append(
                        [rc, cc, rc+int(cell[1]["rowspan"])-1, cc+int(cell[1]["colspan"])-1])
                elif("rowspan" in cell[1]):
                    spanlist.append([rc, cc, rc+int(cell[1]["rowspan"])-1, cc])
                elif("colspan" in cell[1]):
                    spanlist.append([rc, cc, rc, cc+int(cell[1]["colspan"])-1])
        for span in spanlist:
            self.TableMergeCell(span)

    def TableMergeCell(self, list):
        PureCells = self.getPureCellList(list[0])
        #self.table.cell(list[0], PureCells[list[1]]).merge(self.table.cell(list[2], PureCells[list[1]]+list[3]-list[1]))
        self.table.cell(list[0], list[1]).merge(self.table.cell(list[2], list[3]))

    def getPureCellList(self, row):
        PureCells = []
        for c, cell in enumerate(self.table.rows[row].cells):
            if(cell._tc.bottom-cell._tc.top == 1 and cell._tc.right-cell._tc.left == 1):
                PureCells.append(c)
        return PureCells

    def getFirCellList(self, row):
        FirstCell = []
        for c, cell in enumerate(self.table.rows[row].cells):
            if(cell._tc.top == row and cell._tc.left == c):
                FirstCell.append(c)
        return FirstCell

    def TableProperty(self, item, widthInches):
        for rc, row in enumerate(item):
            FirstCell = self.getFirCellList(rc)
            for cc, cell in enumerate(row):
                width = 100
                try:
                    width = float(cell[1]["width"].split(
                        "%")[0])*widthInches/100
                    self.table.cell(rc, FirstCell[cc]).width = Inches(width)
                except:
                    try:
                        width = float(cell[1]["style"].split("width:")[
                            1].split("pt")[0])/4*widthInches/100
                        self.table.cell(
                            rc, FirstCell[cc]).width = Inches(width)
                    except:
                        try:
                            width = float(cell[1]["style"].split("width:")[
                                1].split("%")[0])*widthInches/100
                            self.table.cell(
                                rc, FirstCell[cc]).width = Inches(width)
                        except:
                            None

                if(cell[0] != ""):
                    self.table.cell(rc, FirstCell[cc]).text = cell[0]
                #color
                try:
                    if("background" in cell[1]["style"]):
                        shading = parse_xml(
                            '<w:shd {} w:fill="DDDDDD"/>'.format(docx.oxml.ns.nsdecls('w')))
                        self.table.cell(rc, FirstCell[cc])._tc.get_or_add_tcPr().append(
                            shading)
                        self.table.cell(
                            rc, FirstCell[cc]).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except:
                    None
                #Text

    def getCellSize(self, row, col):
        height = self.table.cell(row, col)._tc.bottom - \
            self.table.cell(row, col)._tc.top
        width = self.table.cell(row, col)._tc.right - \
            self.table.cell(row, col)._tc.left
        #startpoint=(table.cell(row,col)._tc.left==col) and (table.cell(row,col)._tc.top==row)
        #print(row,col,(width,height),startpoint)
        return (width, height)

    def TableAlign(self, position="center"):
        #Table Paragraph Align Center
        tblPr = self.table._tbl.tblPr
        tbljc = OxmlElement('w:jc')
        tbljc.set(qn('w:val'), position)
        tblPr.append(tbljc)

    def TableColumnsWidth(self, sizeInInches):
        for c, size in enumerate(sizeInInches):
            for cell in self.table.columns[c].cells:
                cell.width = Inches(size)

    def TableCellBGColor(self, cells, color="DDDDDD"):
        for r, c in cells:
            shading = parse_xml(
                '<w:shd {} w:fill="DDDDDD"/>'.format(docx.oxml.ns.nsdecls('w')))
            self.table.cell(r, c)._tc.get_or_add_tcPr().append(shading)
            self.table.cell(
                r, c).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def TableCellBackgroundColor(self, RowIndex, ColIndex, color="DDDDDD"):
        shading = parse_xml(
            '<w:shd {} w:fill="DDDDDD"/>'.format(docx.oxml.ns.nsdecls('w')))
        self.table.rows[RowIndex].cells[ColIndex]._tc.get_or_add_tcPr().append(
            shading)
        self.table.rows[RowIndex].cells[ColIndex].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def TableBordersRows(self, rowIndex, size, position="top"):
        if(position == 'top'):
            for cell in self.table.rows[rowIndex].cells:
                tcPr = cell._tc.tcPr
                tcBorders = OxmlElement('w:tcBorders')
                top = OxmlElement('w:'+position)
                top.set(qn('w:val'), 'single')
                top.set(qn('w:sz'), str(size))
                top.set(qn('w:space'), '0')
                top.set(qn('w:color'), 'auto')
                tcBorders.append(top)
                tcPr.append(tcBorders)

        if(position == 'bottom'):
            for cell in self.table.rows[rowIndex].cells:
                tcPr = cell._tc.tcPr
                tcBorders = OxmlElement('w:tcBorders')
                top = OxmlElement('w:bottom')
                top.set(qn('w:val'), 'single')
                top.set(qn('w:sz'), str(size))
                top.set(qn('w:space'), '0')
                top.set(qn('w:color'), 'auto')
                tcBorders.append(top)
                tcPr.append(tcBorders)

    def TableBordersColumns(self, colIndex, size, position="left"):
        if(position == 'left'):
            for cell in self.table.columns[colIndex].cells:
                tcPr = cell._tc.tcPr
                tcBorders = OxmlElement('w:tcBorders')
                top = OxmlElement('w:left')
                top.set(qn('w:val'), 'single')
                top.set(qn('w:sz'), str(size))
                top.set(qn('w:space'), '0')
                top.set(qn('w:color'), 'auto')
                tcBorders.append(top)
                tcPr.append(tcBorders)

        if(position == 'right'):
            for cell in self.table.columns[colIndex].cells:
                tcPr = cell._tc.tcPr
                tcBorders = OxmlElement('w:tcBorders')
                top = OxmlElement('w:right')
                top.set(qn('w:val'), 'single')
                top.set(qn('w:sz'), str(size))
                top.set(qn('w:space'), '0')
                top.set(qn('w:color'), 'auto')
                tcBorders.append(top)
                tcPr.append(tcBorders)

    def TableBordersAllCells(self, size):
        for r in range(len(self.table.rows)):
            for c in range(len(self.table.rows[0].cells)):
                tcPr = self.table.cell(r, c)._tc.tcPr
                tcBorders = OxmlElement('w:tcBorders')

                left = OxmlElement('w:left')
                left.set(qn('w:val'), 'single')
                left.set(qn('w:sz'), str(size))
                left.set(qn('w:space'), '0')
                left.set(qn('w:color'), 'auto')
                tcBorders.append(left)

                right = OxmlElement('w:right')
                right.set(qn('w:val'), 'single')
                right.set(qn('w:sz'), str(size))
                right.set(qn('w:space'), '0')
                right.set(qn('w:color'), 'auto')
                tcBorders.append(right)

                top = OxmlElement('w:top')
                top.set(qn('w:val'), 'single')
                top.set(qn('w:sz'), str(size))
                top.set(qn('w:space'), '0')
                top.set(qn('w:color'), 'auto')
                tcBorders.append(top)

                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), str(size))
                bottom.set(qn('w:space'), '0')
                bottom.set(qn('w:color'), 'auto')
                tcBorders.append(bottom)
                tcPr.append(tcBorders)

    def TableBorders(self, size):
        self.TableBordersAllCells(5)
        self.TableBordersRows(0, size, "top")
        self.TableBordersRows(-1, size, "bottom")
        self.TableBordersColumns(0, size, "left")
        self.TableBordersColumns(-1, size, "right")
